"""
Document Parsers Module
Handles parsing of text, PDF, and Word documents.
"""

import os
from typing import Tuple

def parse_text(content: str) -> str:
    """Parse plain text content."""
    return content.strip()


def parse_text_file(file_path: str) -> str:
    """Parse a plain text file."""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read().strip()


def parse_pdf(file_path: str) -> str:
    """
    Parse a PDF file and extract text content.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n'.join(text_parts).strip()
    
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: str) -> str:
    """
    Parse a Word document and extract text content.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Extracted text from the document
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts).strip()
    
    except Exception as e:
        raise ValueError(f"Failed to parse Word document: {str(e)}")


def detect_and_parse(file_path: str) -> Tuple[str, str]:
    """
    Auto-detect file type and parse accordingly.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Tuple of (extracted_text, file_type)
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return parse_pdf(file_path), 'pdf'
    elif ext == '.docx':
        return parse_docx(file_path), 'docx'
    elif ext in ['.txt', '.md', '.text']:
        return parse_text_file(file_path), 'text'
    else:
        # Try to read as plain text
        try:
            return parse_text_file(file_path), 'text'
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def parse_uploaded_file(file_storage, filename: str) -> Tuple[str, str]:
    """
    Parse an uploaded file from Flask request.
    
    Args:
        file_storage: Flask FileStorage object
        filename: Original filename
        
    Returns:
        Tuple of (extracted_text, file_type)
    """
    import tempfile
    
    ext = os.path.splitext(filename)[1].lower()
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        file_storage.save(tmp.name)
        tmp_path = tmp.name
    
    try:
        text, file_type = detect_and_parse(tmp_path)
        return text, file_type
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
